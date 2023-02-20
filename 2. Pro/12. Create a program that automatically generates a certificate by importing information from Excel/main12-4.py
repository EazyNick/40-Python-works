#docx를 pdf로 바꿔주자.
from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

load_wb = load_workbook(r"40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\수료증명단.xlsx")
load_ws = load_wb.active

name_list = [] #빈 리스트 3개 생성
birthday_list = [] 
ho_list = []
for i in range(1, load_ws.max_row + 1): #엑셀파일의 마지막줄까지 줄을 읽고 어펜드로 리스트값들 추가해줌.
#3가지 리스트 만들어 각 리스트 값들을 어펜드 해주고 3가지 리스트를 출력.
    name_list.append(load_ws.cell(i, 1). value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value) 
#이름 생일 호 를 리스트로 가져옴.

print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)): #사람 수만큼 반복
    doc = docx.Document(r'40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\수료증양식.docx')

    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:esstAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph() #paragraph는 문단을 추가하는 것.
    run = para.add_run('\n\n') #추가할 문장
    run = para.add_run('            제 '+ ho_list[i] + '호\n') #추가할 문장
    run.font.name = '나눔고딕' # 폰트체 나눔고딕
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕') #한글 폰트를 따로 설정해 둔다.
    #이걸 안하면 영어만 폰트되고 한국어는 폰트처리가 되지 않는다.
    run.font.size = docx.shared.Pt(20) #폰트 크기 20

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수  료  증')
    run.font.name = '나눔고딕'
    run.bold = True #해당 런에 다양한 스타일 적용하겠다.
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('      성         명 :' + name_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('      생 년 월 일 :' + birthday_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('      교 육 과 정 : 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('      교 육 날 짜: 2021.08. 05~2021.09.09\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\\'+name_list[i]+'.docx')
    convert('40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\\'+name_list[i]+'.docx',
            '40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\\'+name_list[i]+'.pdf')
    #docx형태로 저장!, 중간에 name_list, birthday_list변수가 포함된거 빼곤 다 12-3과 같다.
    #convert로 pdf로 변환
    #내가 docx가 없어서 생성이 안되는 것 같다.