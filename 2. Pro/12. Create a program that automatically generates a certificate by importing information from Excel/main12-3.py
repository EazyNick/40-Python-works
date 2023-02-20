#수료증 제작
#안에 띄어쓰기 몇번인가 등은 다 직접 알아봐야함.
#파일이 .docx로 저장되있는지 잘 확인하자!
import docx
from docx.oxml.ns import qn
import docx.enum.text
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r'40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\수료증양식.docx')

style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:esstAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph() #paragraph는 문단을 추가하는 것.
run = para.add_run('\n\n') #추가할 문장
run = para.add_run('            제 2020-0001 호\n') #추가할 문장
run.font.name = '나눔고딕' # 폰트체 나눔고딕
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕') #한글 폰트를 따로 설정해 둔다.
#이걸 안하면 영어만 폰트되고 한국어는 폰트처리가 되지 않는다.
run.font.size = docx.shared.Pt(20) #폰트 크기 20

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수  료  증')
run.font.name = '나눔고딕'
run.bold = True #해당 런에 다양한 스타일 적용하겠다.
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('      성         명 : 장다인\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('      생 년 월 일 : 2017.12.12\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('      교 육 과 정 : 파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('      교 육 날 짜: 2021.08. 05~2021.09.09\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

doc.save(r'40 Python works\12. Create a program that automatically generates a certificate by importing information from Excel\수료증양식.docx')
#docx형태로 저장! 난 docx가 없어서 미리 만들어두면 저장만 되고 자동 생성은 안된다..

